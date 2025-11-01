from django.conf import settings
from app_cg.models import Templates
from .Variavel import Variavel
from .ContratosC import ContratosC
from django.core.files.storage import default_storage
from datetime import datetime
from docx import Document
import re

class Template:
    def __init__(self, codtemplate=0, codempresa=0, nome=None, descricao=None, template_url=None, dtatualiz=datetime.now(), status=1):
        self.codtemplate = codtemplate
        self.codempresa = codempresa
        self.nome = nome
        self.descricao = descricao
        self.template_url = template_url
        self.dtatualiz = dtatualiz
        self.status = status

    def obterTemplates(self, codempresa, codtemplate=None):
        if codtemplate is None:
            # Se o codtemplate for None, então retorna uma lista de templates da empresa selecionada
            return list(Templates.objects.filter(codempresa=codempresa, status=1))
        # Se o codtemplate estiver preenchido, então retorna o Objeto template, se não, vai retornar None
        return Templates.objects.filter(codtemplate=codtemplate).first()
    
    def obterArquivoTemplate(self):
        # Sempre usar rb (read binary) para arquivos binários, como .docx. 
        with default_storage.open(self.template_url, "rb") as arquivo:
            return arquivo
    
    def obterArquivoAjuda(self):
        caminho_arquivo_ajuda = '/DOCUMENTACAO/GUIAS/Guia_Template_DocFlow.docx'
        with default_storage.open(caminho_arquivo_ajuda, "rb") as arquivo:
            return arquivo

    def obterInstanciaTemplateCompletoPorCodtemplate(self, codempresa, codtemplate):
        template_obj = self.obterTemplates(codempresa, codtemplate)
        self.obterInstanciaTemplateCompletoPorSemBancoDeDados(template_obj=template_obj)

    def obterInstanciaTemplateCompletoPorSemBancoDeDados(self, template_obj):
        self.codtemplate = template_obj.codtemplate
        self.codempresa = template_obj.codempresa
        self.nome = template_obj.nome
        self.descricao = template_obj.descricao
        self.template_url = template_obj.template_url
        self.dtatualiz = template_obj.dtatualiz
        self.status = template_obj.status
    
    def obterObjetoTemplateSemCodempresa(self, codtemplate):
        return Templates.objects.get(codtemplate=codtemplate)

    def salvarTemplate(self, arquivo_template):
        ## primeiro salva as informações no banco de dados
        template_obj = Templates(codempresa=self.codempresa,nome=self.nome,
                                 descricao=self.descricao, dtatualiz=self.dtatualiz, status=self.status)
        template_obj.save()
        ## guardando o código do template salvo
        self.codtemplate = template_obj.codtemplate 
        ## agora salva o template na nuvem do S3
        caminho = default_storage.save(self.criar_caminho_template(), arquivo_template)
        ## guardar o caminnho do template do S3 no banco de dados
        template_obj.template_url = self.template_url = str(caminho)
        template_obj.save()

    def atualizarTemplate(self, arquivo_template):
        template_obj = self.obterTemplates(self.codempresa, self.codtemplate)
        template_obj.nome = self.nome
        template_obj.descricao = self.descricao
        template_obj.dtatualiz = datetime.now()
        self.template_url = template_obj.template_url
        print(f'\nDEBUG:\n  URL do template: {self.template_url}')
        if arquivo_template and default_storage.exists(self.template_url):
            default_storage.delete(self.template_url) # Excluir arquivo antigo no S3 e cadastrar o novo
            caminho = default_storage.save(self.criar_caminho_template(), arquivo_template)
            template_obj.template_url = self.template_url = str(caminho)
        template_obj.save() # Salvar o template no banco de dados

    def excluirTemplate(self):
        if self.codtemplate != 0: # Verifica se o código do template foi informado
            template_obj = self.obterTemplates(self.codempresa, self.codtemplate) # Obtém o objeto template 
            # Validar se existe contratos vinculados a este templates antes de excluir
            if ContratosC().existeContratoGeradoPeloCodtemplate(self.codtemplate): 
                template_obj.status = 0 # Desativa o template
                template_obj.dtatualiz = datetime.now() # Data de atualização do template (data de desativação)
                template_obj.save() # Salvar objeto no banco de dados
                return 2 # Retorno 2: existe contrato vinculado ao template
            self.obterInstanciaTemplateCompletoPorSemBancoDeDados(template_obj=template_obj) # Preencher os atributos do objeto do template
            default_storage.delete(template_obj.template_url) # Exclui arquivo do template do S3
            template_obj.delete() # Exclui o registro do banco de dados
            return 1 # Retorno 1: Objeto deletado com sucesso
        return 3 # Retorno 3: não foi informado o código do template
    
    def extrair_variaveis(self):
        # Obter o arquivo do template selecionado e abri-lo
        # doc = self.obterArquivoTemplate()
        with default_storage.open(self.template_url, "rb") as arquivo:
            doc = Document(arquivo)
        # Obter o texto completo do documento do template
        # texto_completo = "\n".join([p.text for p in doc.paragraphs])
        texto_completo = extrair_texto(documento=doc)
        # Expressão regular para encontrar padrões como <?tipo:nome:descricao?>
        padrao = re.compile(r"<\?([^:<>]+):([^:<>]+):([^:<>]+)\?>")
        # Extrair apenas os tipos de dados existentes no sistema
        tipos_permitidos = {"palavra", "inteiro", "moeda", "data", "hora"}
        resultado_json = [{"tipo": m.group(1), "nome": m.group(2), "descricao": m.group(3)} 
                          for m in padrao.finditer(texto_completo)
                          if m.group(1) in tipos_permitidos
                          ]
        # print(f'DEBUG: resultado_json = {resultado_json}')
        return resultado_json
    
    def atualizar_variaveis(self):
        v = Variavel(codtemplate=self.codtemplate) # Coletar variável do template
        json = self.extrair_variaveis() # Extrair variáveis do novo template
        if json == []:
            return False
        v.excluirVariavel() # Excluir as variáveis existentes
        nova_variavel = Variavel(codtemplate=self.codtemplate, variaveis=json)
        nova_variavel.salvarVariavel() # Salvar variáveis novamente no banco de dados
        return True

    def criar_caminho_template(self):
        if settings.DEBUG:
            return f'HOMOLOGACAO/templates/empresa_{self.codempresa.codempresa}/{self.codtemplate}.docx'
        return f'PRODUCAO/templates/empresa_{self.codempresa.codempresa}/{self.codtemplate}.docx'
    
def extrair_texto(documento):
    texto = []
    
    def ler_elementos(container):
        for p in container.paragraphs:
            texto.append(p.text)
        for t in container.tables:
            for linha in t.rows:
                for celula in linha.cells:
                    ler_elementos(celula)
    
    ler_elementos(documento)
    return "\n".join(texto)