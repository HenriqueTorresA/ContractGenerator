from app_cg.models import Contratos, V_BuscaContratos, V_ObterDadosContrato
from django.conf import settings
import re, datetime, os, subprocess, platform
from django.core.files.storage import default_storage
from docx import Document
from io import BytesIO
from docx.text.run import Run
from docx.enum.text import WD_BREAK
from .Definicoes import Definicoes
from tempfile import NamedTemporaryFile

class ContratosC:
    def __init__(self, codcontrato=None,codusuario=None,codtemplate=None,codempresa=None,nome_arquivo=None,contrato_url=None,contrato_json=None,status=None,dtatualiz=None):
        self.codcontrato = codcontrato 
        self.codusuario = codusuario 
        self.codtemplate = codtemplate 
        self.codempresa = codempresa
        self.nome_arquivo = nome_arquivo
        self.contrato_url = contrato_url 
        self.contrato_json = contrato_json 
        self.status = status 
        self.dtatualiz = dtatualiz 
    
    def obterContratos(self, codempresa, codcontrato=None):
        if codcontrato is None:
            # Se o codcontrato for None, então retorna uma lista de contratos da empresa selecionada
            return list(Contratos.objects.filter(codempresa=codempresa))
        # Se o codcontrato estiver preenchido, então retorna o Objeto contrato, se não, vai retornar None
        # Realiza também tratativa de filtrar apenas documentos da empresa do usuário logado
        contrato_obj = Contratos.objects.filter(codempresa=codempresa, codcontrato=codcontrato).first()
        if contrato_obj is None: return contrato_obj # Se o contrato não for encontrado, então retorna None
        self.codcontrato = contrato_obj.codcontrato
        self.codusuario = contrato_obj.codusuario
        self.codtemplate = contrato_obj.codtemplate
        self.codempresa = contrato_obj.codempresa
        self.nome_arquivo = contrato_obj.nome_arquivo
        self.contrato_url = contrato_obj.contrato_url
        self.contrato_json = contrato_obj.contrato_json
        self.status = contrato_obj.status
        self.dtatualiz = contrato_obj.dtatualiz
        return contrato_obj
    
    def vBuscaContratos(self, codempresa):
        return list(V_BuscaContratos.objects.filter(codempresa=codempresa))

    def vBuscaDadosContrato(self, codempresa, codcontrato):
        return V_ObterDadosContrato.objects.filter(codempresa=codempresa, codcontrato=codcontrato).first()

    def obterArquivoContrato(self, operacao=2):
        if operacao == '1': # Permite download do arquivo em PDF
            print(f'DEBUG: obtendo arquivo contrato em PDF, operacao {operacao}')
            # Caso contrário, seguir com a conversão do DOCX para PDF
            with default_storage.open(self.contrato_url, "rb") as arquivo:
                docx_bytes = arquivo.read()
            pdf_bytes = docx_para_pdf_bytes(docx_bytes)
            return pdf_bytes
        else: # Permite download do arquivo em DOCX
            print('DEBUG: obtendo arquivo contrato em DOCX')
            with default_storage.open(self.contrato_url, "rb") as arquivo:
                return arquivo

    def gerarContrato(self):
        # Capturar o template
        with default_storage.open(self.codtemplate.template_url, "rb") as arquivo:
            doc = Document(arquivo)

        # Regex para identificar as expressões no formato <?tipo:nome:descricao?>
        padrao = r"<\?([a-zA-Z0-9_]+):([a-zA-Z0-9_]+):.*?\?>"

        doc = substituir_variaveis_docx(doc, self.contrato_json)
                    
        # # CONVERTER PARA PDF COM LIBREOFFICE --> Isso não funciona na Vercel, precisaria de um servidor como EC2 ou Docker
        # subprocess.run([
        #     "soffice", 
        #     "--headless",  # Sem interface gráfica
        #     "--convert-to", 
        #     "pdf",  # Converter para PDF
        #     modified_docx_path
        # ], check=True)

        # Salvar o DOCX gerado em memória para o S3
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        # Salvar os dados do contrato no banco de dados
        contrato_obj = Contratos(codusuario=self.codusuario,codtemplate=self.codtemplate,codempresa=self.codempresa,
                                 contrato_json=self.contrato_json,nome_arquivo=self.nome_arquivo,status=self.status,
                                 dtatualiz=self.dtatualiz)
        contrato_obj.save() # Salvar o objeto contrato no banco de dados
        self.codcontrato = contrato_obj.codcontrato # Coletar o código do contrato que acabou de ser salvo
        caminho = default_storage.save(self.criarCaminhoContrato(), buffer) # Salvar o arquivo na nuvem do S3
        contrato_obj.contrato_url = str(caminho) # Guardar o caminho do S3 no banco de dados
        contrato_obj.save() # Atualizar o registro do caminho do contrato no banco de dados
        print(f'DEBUG: contrato salvo!\ncaminho: {caminho}')
        if caminho: 
            return True 
        return False # Se não conseguiu salvar o arquivo, retorna False e mostra erro na tela do usuário

    def criarCaminhoContrato(self): # Obtém o caminho do novo contrato na nuvem do S3
        caminho_inicial = 'HOMOLOGACAO' if settings.DEBUG else 'PRODUCAO'
        caminho_final = f'contratos/empresa_{self.codusuario.codempresa.codempresa}/{self.codcontrato}.docx'
        return f'{caminho_inicial}/{caminho_final}'
    
    def existeContratoGeradoPeloCodtemplate(self, codtemplate): # Verifica se existe contrato gerado a partir do código do template informado
        contrato_obj = Contratos.objects.filter(codtemplate=codtemplate).first()
        if contrato_obj is None:
            return False
        return True
    
    def excluirContrato(self):
        if self.codcontrato != 0: # Verifica se o código do contrato foi informado
            contrato_obj = self.obterContratos(self.codempresa, self.codcontrato) # Obtém o objeto contrato 
            default_storage.delete(contrato_obj.contrato_url) # Exclui arquivo do contrato do S3
            contrato_obj.delete() # Exclui o registro do banco de dados
            return 1 # Retorno 1: Objeto deletado com sucesso
        return 3 # Retorno 3: não foi informado o código do contrato

def ifnull(x, y):
    if x is not None and x != "" and x != " ":
        return x
    else:
        return y
    
def transformaData(data_str):
    try:
        data = datetime.datetime.strptime(data_str, "%Y-%m-%d")
    except ValueError:
        return None
    meses = [
        "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
        "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"
    ]
    
    return f"{data.day} de {meses[data.month - 1]} de {data.year}"

def trataNomeVariavel(nome):
    nome = str(nome)
    return nome.capitalize()

def substituir_variaveis_docx(doc, contrato_json):
    padrao = re.compile(r"<\?([^:<>]+):([^:<>]+):[^<>]*\?>")  # <?tipo:nome:descricao?>
    
    for p in doc.paragraphs:
        _processar_paragrafo(p, padrao, contrato_json)

    # Percorrer todas as tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    _processar_paragrafo(p, padrao, contrato_json)
    
    return doc


def _processar_paragrafo(p, padrao, contrato_json):
    # lista de runs e seus textos
    runs = list(p.runs)
    if not runs:
        return

    texts = [r.text for r in runs]
    combined = "".join(texts)

    # encontra todos os matches no texto combinado
    matches = list(re.finditer(padrao, combined))
    if not matches:
        return

    # Prepara limites cumulativos para mapear índices do texto combinado -> runs
    cum_limits = []
    offset = 0
    for t in texts:
        cum_limits.append((offset, offset + len(t)))  # (start_in_combined, end_in_combined)
        offset += len(t)

    # Processa matches de trás pra frente para não invalidar índices
    for m in reversed(matches):
        tipo = m.group(1)
        nome = m.group(2)
        start_idx, end_idx = m.start(), m.end()

        nome_formatado = trataNomeVariavel(nome)

        # validações
        if nome_formatado not in contrato_json.get("dados_json", {}):
            continue
        if tipo not in Definicoes.TIPOS_PERMITIDOS:
            continue

        dado = contrato_json["dados_json"].get(nome_formatado)
        if tipo == "data":
            valor = ifnull(transformaData(dado), "___ / ___ / _____")
        elif tipo == "listaenumerada":
            if isinstance(dado, list) and dado:
                lista_itens = [item for sub in dado for item in (sub if isinstance(sub, list) else [sub])]
                texto_paragrafo = ''.join(run.text for run in p.runs)
                padrao = rf"<\?{tipo}:{nome}:.*?\?>"
                match = re.search(padrao, texto_paragrafo)
                # Substituir o placeholder por uma lista numerada com quebras reais
                
                if match:
                    # Limpa todo o parágrafo
                    for run in p.runs:
                        run.text = ""
                    # Adiciona a lista formatada
                    for idx, item in enumerate(lista_itens):
                        new_run = p.add_run(f"            {idx + 1}. {item}")
                        new_run.add_break(WD_BREAK.LINE)
                    continue
            else:
                valor = ''
        elif tipo == "palavrasemlinha":
            valor = ifnull(dado, '')
        else:
            valor = ifnull(dado, '______________________________')

        # encontra run inicial e run final que cobrem start_idx e end_idx-1
        first_run_idx = None
        last_run_idx = None
        start_offset_in_first = None
        end_offset_in_last = None

        for i, (a, b) in enumerate(cum_limits):
            if a <= start_idx < b:
                first_run_idx = i
                start_offset_in_first = start_idx - a
            if a <= end_idx - 1 < b:
                last_run_idx = i
                end_offset_in_last = end_idx - a
            if first_run_idx is not None and last_run_idx is not None:
                break

        # se não conseguimos mapear corretamente, pular
        if first_run_idx is None or last_run_idx is None:
            continue

        # prepara novos textos dos runs afetados
        # parte anterior ao placeholder no primeiro run
        prefix = texts[first_run_idx][:start_offset_in_first]
        # parte posterior ao placeholder no último run
        suffix = texts[last_run_idx][end_offset_in_last:]

        # novo texto que ficará no primeiro run
        novo_texto_primeiro_run = prefix + valor + suffix

        # aplica ao primeiro run (preserva a formatação desse run)
        runs[first_run_idx].text = novo_texto_primeiro_run
        texts[first_run_idx] = novo_texto_primeiro_run

        # remove os runs intermediários (do last down to first+1)
        # usamos p._element.remove para remover do xml do parágrafo
        # e atualizamos as listas texts/runs/cum_limits
        if last_run_idx > first_run_idx:
            # remover do último para o primeiro+1 para não mexer nos índices
            for rem_idx in range(last_run_idx, first_run_idx, -1):
                try:
                    p._element.remove(runs[rem_idx]._element)
                except Exception:
                    # em caso de erro, apenas passe (não fatal)
                    pass
                del runs[rem_idx]
                del texts[rem_idx]
                del cum_limits[rem_idx]

            # após remoção, recalcula cum_limits (opcional mas seguro)
            # recomputar cum_limits para próximos matches
            new_cum = []
            off = 0
            for t in texts:
                new_cum.append((off, off + len(t)))
                off += len(t)
            cum_limits = new_cum

def docx_para_pdf_bytes(docx_bytes):
    soffice_cmd = get_soffice_command()

    with NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
        temp_docx.write(docx_bytes)
        temp_docx_path = temp_docx.name

    temp_pdf_path = temp_docx_path.replace(".docx", ".pdf")

    subprocess.run([
        soffice_cmd,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", os.path.dirname(temp_docx_path),
        temp_docx_path
    ], check=True)

    # Lê o PDF gerado
    with open(temp_pdf_path, "rb") as f:
        pdf_bytes = f.read()

    # Remove temporários
    os.remove(temp_docx_path)
    os.remove(temp_pdf_path)

    return pdf_bytes

def get_soffice_command():
    system = platform.system()

    if system == "Windows":
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
        ]
        for p in possible_paths:
            if os.path.exists(p):
                return p
        raise FileNotFoundError("LibreOffice não encontrado no Windows.")
    else:
        # Linux
        return "libreoffice"   # ou "soffice"